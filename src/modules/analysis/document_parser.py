"""
Document Parser Sub-Agent

Specialized sub-agent for parsing and extracting content from various document formats.
Handles PDF, DOCX, TXT, and Markdown files with intelligent content structuring.
"""

import logging
import os
from typing import Dict, Any, List, Optional
from pathlib import Path
from datetime import datetime
import asyncio

from ...agents.base_agent import BaseAgent

logger = logging.getLogger(__name__)


class DocumentParser(BaseAgent):
    """Sub-agent for document parsing and content extraction."""
    
    def __init__(self):
        super().__init__(
            name="DocumentParser", 
            description="Extracts and structures content from various document formats"
        )
        self.supported_formats = ['.txt', '.md', '.pdf', '.docx']
        self.extraction_stats = {
            'documents_processed': 0,
            'pages_processed': 0,
            'file_size': 0,
            'file_extension': '',
            'total_pages': 0,
            'avg_processing_time': 0.0
        }
    
    async def process(self, input_data: Dict[str, Any]) -> Dict[str, Any]:
        """
        Main entry point for document processing.
        
        Args:
            input_data: Dictionary containing 'document_path' and optional parameters
            
        Returns:
            Dictionary with processing results
        """
        return await self.parse_document(input_data)
    
    async def parse_document(self, input_data: Dict[str, Any]) -> Dict[str, Any]:
        """Parse document and extract content with metadata."""
        document_path = input_data.get('document_path', 'unknown')
        
        try:
            # Validate input data
            if not input_data or not isinstance(input_data, dict):
                raise ValueError("Invalid input data: must be a non-empty dictionary")
            
            if not document_path or document_path == 'unknown':
                raise ValueError("Document path is required and cannot be empty")
            
            file_path = Path(document_path)
            
            # Comprehensive file validation
            if not file_path.exists():
                raise FileNotFoundError(f"Document not found: {file_path}")
            
            if not file_path.is_file():
                raise ValueError(f"Path is not a file: {file_path}")
            
            # Check file permissions
            if not os.access(file_path, os.R_OK):
                raise PermissionError(f"No read permission for file: {file_path}")
            
            self.logger.info(f"Parsing document: {file_path}")
            
            # Get file info with error handling
            try:
                file_stat = file_path.stat()
                file_size = file_stat.st_size
                file_extension = file_path.suffix.lower()
                
                # Check for empty files
                if file_size == 0:
                    self.logger.warning(f"File is empty: {file_path}")
                    return {
                        'status': 'warning',
                        'warning': 'File is empty',
                        'document_path': str(file_path),
                        'content': '',
                        'structured_data': {},
                        'metadata': {
                            'file_size': 0,
                            'file_extension': file_extension,
                            'processing_timestamp': self._get_timestamp()
                        },
                        'extraction_stats': self.extraction_stats.copy()
                    }
                
                # File size limits (100MB max)
                max_file_size = 100 * 1024 * 1024  # 100MB
                if file_size > max_file_size:
                    raise ValueError(f"File too large: {file_size} bytes (max: {max_file_size})")
                
            except OSError as e:
                raise IOError(f"Cannot access file metadata: {e}")
            
            self.extraction_stats['file_size'] = file_size
            self.extraction_stats['file_extension'] = file_extension
            
            # Extract content with enhanced error handling
            try:
                content = await self._extract_content(file_path, file_extension)
                
                if not content or not content.strip():
                    self.logger.warning(f"No content extracted from: {file_path}")
                    content = ""
                    
            except Exception as e:
                raise RuntimeError(f"Content extraction failed: {e}")
            
            # Parse and structure content with error handling
            try:
                structured_data = await self._parse_content(content) if content else {}
            except Exception as e:
                self.logger.warning(f"Content parsing failed, using raw content: {e}")
                structured_data = {'raw_content': content}
            
            result = {
                'status': 'success',
                'document_path': str(file_path),
                'content': content,
                'structured_data': structured_data,
                'metadata': {
                    'file_size': file_size,
                    'file_extension': file_extension,
                    'processing_timestamp': self._get_timestamp()
                },
                'extraction_stats': self.extraction_stats.copy()
            }
            
            self.log_operation("Document parsing completed", {'document': str(file_path)})
            return result
            
        except (ValueError, FileNotFoundError, PermissionError, RuntimeError) as e:
            # Expected errors - log and return structured error response
            error_msg = str(e)
            self.logger.error(f"Document parsing failed: {error_msg}")
            return {
                'status': 'error',
                'error': error_msg,
                'error_type': type(e).__name__,
                'document_path': document_path
            }
            
        except Exception as e:
            # Unexpected errors - log with more detail
            error_msg = f"Unexpected error during document parsing: {str(e)}"
            self.logger.error(error_msg, exc_info=True)
            return {
                'status': 'error',
                'error': error_msg,
                'error_type': 'UnexpectedError',
                'document_path': document_path
            }
    
    async def _extract_content(self, file_path: Path, file_extension: str) -> str:
        """Extract raw content from document based on file type."""
        try:
            if file_extension == '.txt':
                return await self._extract_txt_content(file_path)
            elif file_extension == '.md':
                return await self._extract_md_content(file_path)
            elif file_extension == '.pdf':
                return await self._extract_pdf_content(str(file_path))
            elif file_extension == '.docx':
                return await self._extract_docx_content(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_extension}")
                
        except Exception as e:
            self.logger.error(f"Content extraction failed: {e}")
            raise
    
    async def _extract_txt_content(self, file_path: Path) -> str:
        """Extract content from text file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except UnicodeDecodeError:
            # Try with different encoding
            with open(file_path, 'r', encoding='latin-1') as file:
                return file.read()
    
    async def _extract_md_content(self, file_path: Path) -> str:
        """Extract content from Markdown file."""
        return await self._extract_txt_content(file_path)
    
    async def _extract_pdf_content(self, file_path: str) -> str:
        """Extract content from PDF file with comprehensive error handling."""
        try:
            # Check if file exists and is readable
            if not Path(file_path).exists():
                raise FileNotFoundError(f"PDF file not found: {file_path}")
            
            if not Path(file_path).is_file():
                raise ValueError(f"Path is not a file: {file_path}")
            
            # Check file size (avoid extremely large files)
            file_size = Path(file_path).stat().st_size
            if file_size > 50 * 1024 * 1024:  # 50MB limit
                raise ValueError(f"PDF file too large: {file_size / (1024*1024):.1f}MB (max 50MB)")
            
            if file_size == 0:
                raise ValueError("PDF file is empty")
            
            # Import PyPDF2 with fallback
            try:
                import PyPDF2
            except ImportError as e:
                self.logger.error("PyPDF2 not installed. Install with: pip install PyPDF2")
                raise ImportError("PyPDF2 dependency required for PDF processing") from e
            
            # Extract text from PDF
            extracted_text = ""
            page_count = 0
            
            with open(file_path, 'rb') as file:
                try:
                    reader = PyPDF2.PdfReader(file)
                    
                    # Check if PDF is encrypted
                    if reader.is_encrypted:
                        self.logger.warning(f"PDF is encrypted: {file_path}")
                        return "[ERROR] PDF is password protected and cannot be processed"
                    
                    # Check if PDF has pages
                    if len(reader.pages) == 0:
                        raise ValueError("PDF contains no pages")
                    
                    # Extract text from each page
                    for page_num, page in enumerate(reader.pages):
                        try:
                            page_text = page.extract_text()
                            if page_text:
                                extracted_text += f"\n--- Page {page_num + 1} ---\n"
                                extracted_text += page_text
                                page_count += 1
                        except Exception as page_error:
                            self.logger.warning(f"Failed to extract text from page {page_num + 1}: {page_error}")
                            continue
                    
                    # Validate extracted content
                    if not extracted_text.strip():
                        return "[WARNING] PDF appears to contain no extractable text (possibly image-based or corrupted)"
                    
                    if len(extracted_text.strip()) < 10:
                        return f"[WARNING] Very little text extracted from PDF: {extracted_text.strip()}"
                    
                    # Log successful extraction
                    self.logger.info(f"Successfully extracted text from {page_count} pages of {file_path}")
                    return extracted_text.strip()
                    
                except PyPDF2.errors.PdfReadError as pdf_error:
                    error_msg = f"Corrupted or invalid PDF file: {pdf_error}"
                    self.logger.error(error_msg)
                    return f"[ERROR] {error_msg}"
                    
                except Exception as read_error:
                    error_msg = f"Failed to read PDF structure: {read_error}"
                    self.logger.error(error_msg)
                    return f"[ERROR] {error_msg}"
            
        except FileNotFoundError as e:
            self.logger.error(f"PDF file not found: {e}")
            return f"[ERROR] File not found: {file_path}"
            
        except ValueError as e:
            self.logger.error(f"PDF validation error: {e}")
            return f"[ERROR] {str(e)}"
            
        except PermissionError as e:
            self.logger.error(f"Permission denied accessing PDF: {e}")
            return f"[ERROR] Permission denied: {file_path}"
            
        except Exception as e:
            error_msg = f"Unexpected error during PDF extraction: {e}"
            self.logger.error(error_msg)
            return f"[ERROR] {error_msg}"
    
    async def _extract_docx_content(self, file_path: Path) -> str:
        """Extract content from DOCX file with comprehensive error handling."""
        try:
            # Convert to Path object if needed
            if isinstance(file_path, str):
                file_path = Path(file_path)
            
            # Check if file exists and is readable
            if not file_path.exists():
                raise FileNotFoundError(f"DOCX file not found: {file_path}")
            
            if not file_path.is_file():
                raise ValueError(f"Path is not a file: {file_path}")
            
            # Check file size (avoid extremely large files)
            file_size = file_path.stat().st_size
            if file_size > 50 * 1024 * 1024:  # 50MB limit
                raise ValueError(f"DOCX file too large: {file_size / (1024*1024):.1f}MB (max 50MB)")
            
            if file_size == 0:
                raise ValueError("DOCX file is empty")
            
            # Import python-docx with fallback
            try:
                from docx import Document
                from docx.opc.exceptions import PackageNotFoundError
            except ImportError as e:
                self.logger.error("python-docx not installed. Install with: pip install python-docx")
                raise ImportError("python-docx dependency required for DOCX processing") from e
            
            # Extract text from DOCX
            extracted_text = ""
            paragraph_count = 0
            table_count = 0
            
            try:
                # Open and read the document
                doc = Document(file_path)
                
                # Extract text from paragraphs
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        extracted_text += paragraph.text + "\n"
                        paragraph_count += 1
                
                # Extract text from tables
                for table in doc.tables:
                    table_count += 1
                    extracted_text += f"\n--- Table {table_count} ---\n"
                    
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            cell_text = cell.text.strip()
                            if cell_text:
                                row_text.append(cell_text)
                        
                        if row_text:
                            extracted_text += " | ".join(row_text) + "\n"
                
                # Extract text from headers and footers
                for section in doc.sections:
                    # Header text
                    if section.header:
                        for paragraph in section.header.paragraphs:
                            if paragraph.text.strip():
                                extracted_text += f"[HEADER] {paragraph.text}\n"
                    
                    # Footer text
                    if section.footer:
                        for paragraph in section.footer.paragraphs:
                            if paragraph.text.strip():
                                extracted_text += f"[FOOTER] {paragraph.text}\n"
                
                # Validate extracted content
                if not extracted_text.strip():
                    return "[WARNING] DOCX appears to contain no extractable text (possibly empty or corrupted)"
                
                if len(extracted_text.strip()) < 10:
                    return f"[WARNING] Very little text extracted from DOCX: {extracted_text.strip()}"
                
                # Log successful extraction
                self.logger.info(f"Successfully extracted text from DOCX: {paragraph_count} paragraphs, {table_count} tables")
                return extracted_text.strip()
                
            except PackageNotFoundError as package_error:
                error_msg = f"Invalid DOCX file format: {package_error}"
                self.logger.error(error_msg)
                return f"[ERROR] {error_msg}"
                
            except Exception as docx_error:
                error_msg = f"Failed to process DOCX content: {docx_error}"
                self.logger.error(error_msg)
                return f"[ERROR] {error_msg}"
            
        except FileNotFoundError as e:
            self.logger.error(f"DOCX file not found: {e}")
            return f"[ERROR] File not found: {file_path}"
            
        except ValueError as e:
            self.logger.error(f"DOCX validation error: {e}")
            return f"[ERROR] {str(e)}"
            
        except PermissionError as e:
            self.logger.error(f"Permission denied accessing DOCX: {e}")
            return f"[ERROR] Permission denied: {file_path}"
            
        except Exception as e:
            error_msg = f"Unexpected error during DOCX extraction: {e}"
            self.logger.error(error_msg)
            return f"[ERROR] {error_msg}"
    
    async def _extract_metadata(self, file_path: Path) -> Dict[str, Any]:
        """Extract metadata from document."""
        try:
            stat = file_path.stat()
            return {
                'filename': file_path.name,
                'file_size': stat.st_size,
                'created_time': stat.st_ctime,
                'modified_time': stat.st_mtime,
                'file_extension': file_path.suffix.lower()
            }
        except Exception as e:
            self.logger.error(f"Metadata extraction failed: {e}")
            return {}
    
    async def _structure_content(self, content: str) -> Dict[str, Any]:
        """Structure the extracted content into logical sections."""
        try:
            # Simple content structuring
            lines = content.split('\n')
            
            # Identify sections based on patterns
            sections = []
            current_section = {'title': 'Introduction', 'content': []}
            
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                
                # Check if line is a potential header
                if (line.isupper() and len(line) < 100) or \
                   (line.startswith('#') and line.count('#') <= 3):
                    # Save current section if it has content
                    if current_section['content']:
                        sections.append(current_section)
                    
                    # Start new section
                    title = line.replace('#', '').strip()
                    current_section = {'title': title, 'content': []}
                else:
                    current_section['content'].append(line)
            
            # Add the last section
            if current_section['content']:
                sections.append(current_section)
            
            return {
                'sections': sections,
                'total_sections': len(sections),
                'total_lines': len(lines),
                'word_count': len(content.split())
            }
            
        except Exception as e:
            self.logger.error(f"Content structuring failed: {e}")
            return {'sections': [], 'error': str(e)}
    
    async def _parse_content(self, content: str) -> Dict[str, Any]:
        """Parse and structure document content."""
        if not content or not content.strip():
            return {
                'sections': [],
                'total_sections': 0,
                'total_lines': 0,
                'word_count': 0
            }
        
        try:
            # Use the existing _structure_content method
            return await self._structure_content(content)
        except Exception as e:
            self.logger.warning(f"Content parsing failed, returning basic structure: {e}")
            return {
                'raw_content': content,
                'word_count': len(content.split()),
                'character_count': len(content),
                'lines': content.count('\n') + 1,
                'error': str(e)
            }
    
    def get_supported_formats(self) -> List[str]:
        """Get list of supported document formats."""
        return self.supported_formats.copy()
    
    def get_statistics(self) -> Dict[str, Any]:
        """Get parsing statistics."""
        return self.extraction_stats.copy()

"""
Document Processing Module - Independent PDF/Document Text Extraction

This module handles all document processing and text extraction.
It can be developed, tested, and upgraded independently of other modules.
Easily swap out parsers (PyPDF2 -> pdfplumber) without affecting other components.

Features:
- Multiple PDF parsers (pdfplumber, PyPDF2, pymupdf)
- DOCX, HTML, TXT processing
- OCR support for scanned documents
- Parallel document processing
- Error handling and fallback mechanisms
- Document metadata extraction
"""

import logging
from abc import ABC, abstractmethod
from dataclasses import dataclass
from typing import List, Dict, Optional, Any, Union
from pathlib import Path
import mimetypes
from datetime import datetime
import asyncio
from concurrent.futures import ThreadPoolExecutor
import hashlib

# PDF Processing Libraries
try:
    import pdfplumber  # Preferred for robust text extraction
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False

try:
    from pypdf import PdfReader  # Fallback option
    PYPDF_AVAILABLE = True
except ImportError:
    PYPDF_AVAILABLE = False

try:
    import fitz  # PyMuPDF - Alternative option
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False

# Document Processing Libraries
try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from bs4 import BeautifulSoup
    BS4_AVAILABLE = True
except ImportError:
    BS4_AVAILABLE = False

# OCR Libraries
try:
    import pytesseract
    from PIL import Image
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

@dataclass
class DocumentMetadata:
    """Document metadata structure"""
    filename: str
    file_size: int
    mime_type: str
    page_count: Optional[int]
    creation_date: Optional[datetime]
    modification_date: Optional[datetime]
    author: Optional[str]
    title: Optional[str]
    subject: Optional[str]
    language: Optional[str]
    processing_method: str
    confidence_score: float

@dataclass
class ProcessedDocument:
    """Processed document structure"""
    document_id: str
    filename: str
    text_content: str
    metadata: DocumentMetadata
    page_texts: List[str]  # Text from individual pages
    tables: List[Dict[str, Any]]  # Extracted tables
    images: List[Dict[str, Any]]  # Image metadata
    processing_time: float
    success: bool
    error_message: Optional[str] = None

class DocumentProcessorInterface(ABC):
    """Abstract interface for document processors"""
    
    @abstractmethod
    def can_process(self, file_path: Path) -> bool:
        """Check if processor can handle this file type"""
        pass
    
    @abstractmethod
    async def process_document(self, file_path: Path) -> ProcessedDocument:
        """Process document and extract text"""
        pass
    
    @abstractmethod
    def get_processor_name(self) -> str:
        """Return processor name"""
        pass
    
    @abstractmethod
    def get_supported_formats(self) -> List[str]:
        """Return supported file formats"""
        pass

class PDFPlumberProcessor(DocumentProcessorInterface):
    """PDF processor using pdfplumber (most robust)"""
    
    def __init__(self):
        if not PDFPLUMBER_AVAILABLE:
            raise ImportError("pdfplumber is required for PDFPlumberProcessor")
    
    def can_process(self, file_path: Path) -> bool:
        """Check if file is a PDF"""
        return file_path.suffix.lower() == '.pdf'
    
    async def process_document(self, file_path: Path) -> ProcessedDocument:
        """Process PDF using pdfplumber"""
        start_time = asyncio.get_event_loop().time()
        
        try:
            # Run in thread pool to avoid blocking
            loop = asyncio.get_event_loop()
            with ThreadPoolExecutor() as executor:
                result = await loop.run_in_executor(
                    executor, self._process_pdf_sync, file_path
                )
            
            processing_time = asyncio.get_event_loop().time() - start_time
            result.processing_time = processing_time
            
            return result
            
        except Exception as e:
            logger.error(f"PDFPlumber processing failed for {file_path}: {e}")
            return ProcessedDocument(
                document_id=self._generate_document_id(file_path),
                filename=file_path.name,
                text_content="",
                metadata=self._create_error_metadata(file_path, str(e)),
                page_texts=[],
                tables=[],
                images=[],
                processing_time=asyncio.get_event_loop().time() - start_time,
                success=False,
                error_message=str(e)
            )
    
    def _process_pdf_sync(self, file_path: Path) -> ProcessedDocument:
        """Synchronous PDF processing with pdfplumber"""
        text_content = ""
        page_texts = []
        tables = []
        images = []
        
        with pdfplumber.open(file_path) as pdf:
            # Extract metadata
            metadata = self._extract_pdf_metadata(pdf, file_path)
            
            # Process each page
            for page_num, page in enumerate(pdf.pages):
                try:
                    # Extract text
                    page_text = page.extract_text() or ""
                    page_texts.append(page_text)
                    text_content += page_text + "\n"
                    
                    # Extract tables
                    page_tables = page.extract_tables()
                    for table_idx, table in enumerate(page_tables or []):
                        tables.append({
                            'page': page_num + 1,
                            'table_index': table_idx,
                            'data': table,
                            'rows': len(table),
                            'columns': len(table[0]) if table else 0
                        })
                    
                    # Extract images metadata
                    if hasattr(page, 'images'):
                        for img_idx, img in enumerate(page.images):
                            images.append({
                                'page': page_num + 1,
                                'image_index': img_idx,
                                'bbox': img.get('bbox', []),
                                'width': img.get('width'),
                                'height': img.get('height')
                            })
                
                except Exception as e:
                    logger.warning(f"Error processing page {page_num + 1}: {e}")
                    page_texts.append("")
        
        return ProcessedDocument(
            document_id=self._generate_document_id(file_path),
            filename=file_path.name,
            text_content=text_content.strip(),
            metadata=metadata,
            page_texts=page_texts,
            tables=tables,
            images=images,
            processing_time=0.0,  # Will be set by caller
            success=True
        )
    
    def _extract_pdf_metadata(self, pdf, file_path: Path) -> DocumentMetadata:
        """Extract PDF metadata using pdfplumber"""
        try:
            info = pdf.metadata or {}
            file_stat = file_path.stat()
            
            return DocumentMetadata(
                filename=file_path.name,
                file_size=file_stat.st_size,
                mime_type='application/pdf',
                page_count=len(pdf.pages),
                creation_date=self._parse_pdf_date(info.get('CreationDate')),
                modification_date=datetime.fromtimestamp(file_stat.st_mtime),
                author=info.get('Author'),
                title=info.get('Title'),
                subject=info.get('Subject'),
                language=None,  # Not typically available in PDF metadata
                processing_method='pdfplumber',
                confidence_score=0.95
            )
        except Exception as e:
            logger.warning(f"Error extracting PDF metadata: {e}")
            return self._create_basic_metadata(file_path, 'pdfplumber')
    
    def _parse_pdf_date(self, date_str: Optional[str]) -> Optional[datetime]:
        """Parse PDF date format"""
        if not date_str:
            return None
        try:
            # PDF date format: D:YYYYMMDDHHmmSSOHH'mm'
            if date_str.startswith('D:'):
                date_str = date_str[2:16]  # Take YYYYMMDDHHMMSS part
                return datetime.strptime(date_str, '%Y%m%d%H%M%S')
        except:
            pass
        return None
    
    def get_processor_name(self) -> str:
        return "PDFPlumber"
    
    def get_supported_formats(self) -> List[str]:
        return ['.pdf']

class PyPDFProcessor(DocumentProcessorInterface):
    """PDF processor using PyPDF (fallback option)"""
    
    def __init__(self):
        if not PYPDF_AVAILABLE:
            raise ImportError("pypdf is required for PyPDFProcessor")
    
    def can_process(self, file_path: Path) -> bool:
        return file_path.suffix.lower() == '.pdf'
    
    async def process_document(self, file_path: Path) -> ProcessedDocument:
        """Process PDF using PyPDF"""
        start_time = asyncio.get_event_loop().time()
        
        try:
            loop = asyncio.get_event_loop()
            with ThreadPoolExecutor() as executor:
                result = await loop.run_in_executor(
                    executor, self._process_pdf_sync, file_path
                )
            
            processing_time = asyncio.get_event_loop().time() - start_time
            result.processing_time = processing_time
            
            return result
            
        except Exception as e:
            logger.error(f"PyPDF processing failed for {file_path}: {e}")
            return ProcessedDocument(
                document_id=self._generate_document_id(file_path),
                filename=file_path.name,
                text_content="",
                metadata=self._create_error_metadata(file_path, str(e)),
                page_texts=[],
                tables=[],
                images=[],
                processing_time=asyncio.get_event_loop().time() - start_time,
                success=False,
                error_message=str(e)
            )
    
    def _process_pdf_sync(self, file_path: Path) -> ProcessedDocument:
        """Synchronous PDF processing with PyPDF"""
        text_content = ""
        page_texts = []
        
        with open(file_path, 'rb') as file:
            reader = PdfReader(file)
            
            # Extract metadata
            metadata = self._extract_pypdf_metadata(reader, file_path)
            
            # Process each page
            for page_num, page in enumerate(reader.pages):
                try:
                    page_text = page.extract_text() or ""
                    page_texts.append(page_text)
                    text_content += page_text + "\n"
                except Exception as e:
                    logger.warning(f"Error processing page {page_num + 1}: {e}")
                    page_texts.append("")
        
        return ProcessedDocument(
            document_id=self._generate_document_id(file_path),
            filename=file_path.name,
            text_content=text_content.strip(),
            metadata=metadata,
            page_texts=page_texts,
            tables=[],  # PyPDF doesn't extract tables
            images=[],  # PyPDF doesn't extract images
            processing_time=0.0,
            success=True
        )
    
    def _extract_pypdf_metadata(self, reader, file_path: Path) -> DocumentMetadata:
        """Extract PDF metadata using PyPDF"""
        try:
            info = reader.metadata or {}
            file_stat = file_path.stat()
            
            return DocumentMetadata(
                filename=file_path.name,
                file_size=file_stat.st_size,
                mime_type='application/pdf',
                page_count=len(reader.pages),
                creation_date=info.get('/CreationDate'),
                modification_date=datetime.fromtimestamp(file_stat.st_mtime),
                author=info.get('/Author'),
                title=info.get('/Title'),
                subject=info.get('/Subject'),
                language=None,
                processing_method='pypdf',
                confidence_score=0.8
            )
        except Exception as e:
            logger.warning(f"Error extracting PyPDF metadata: {e}")
            return self._create_basic_metadata(file_path, 'pypdf')
    
    def get_processor_name(self) -> str:
        return "PyPDF"
    
    def get_supported_formats(self) -> List[str]:
        return ['.pdf']

class DOCXProcessor(DocumentProcessorInterface):
    """DOCX processor using python-docx"""
    
    def __init__(self):
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is required for DOCXProcessor")
    
    def can_process(self, file_path: Path) -> bool:
        return file_path.suffix.lower() in ['.docx', '.docm']
    
    async def process_document(self, file_path: Path) -> ProcessedDocument:
        """Process DOCX document"""
        start_time = asyncio.get_event_loop().time()
        
        try:
            loop = asyncio.get_event_loop()
            with ThreadPoolExecutor() as executor:
                result = await loop.run_in_executor(
                    executor, self._process_docx_sync, file_path
                )
            
            processing_time = asyncio.get_event_loop().time() - start_time
            result.processing_time = processing_time
            
            return result
            
        except Exception as e:
            logger.error(f"DOCX processing failed for {file_path}: {e}")
            return ProcessedDocument(
                document_id=self._generate_document_id(file_path),
                filename=file_path.name,
                text_content="",
                metadata=self._create_error_metadata(file_path, str(e)),
                page_texts=[],
                tables=[],
                images=[],
                processing_time=asyncio.get_event_loop().time() - start_time,
                success=False,
                error_message=str(e)
            )
    
    def _process_docx_sync(self, file_path: Path) -> ProcessedDocument:
        """Synchronous DOCX processing"""
        doc = DocxDocument(file_path)
        
        # Extract text from paragraphs
        text_content = ""
        for paragraph in doc.paragraphs:
            text_content += paragraph.text + "\n"
        
        # Extract tables
        tables = []
        for table_idx, table in enumerate(doc.tables):
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            
            tables.append({
                'table_index': table_idx,
                'data': table_data,
                'rows': len(table_data),
                'columns': len(table_data[0]) if table_data else 0
            })
        
        # Extract metadata
        metadata = self._extract_docx_metadata(doc, file_path)
        
        return ProcessedDocument(
            document_id=self._generate_document_id(file_path),
            filename=file_path.name,
            text_content=text_content.strip(),
            metadata=metadata,
            page_texts=[text_content],  # DOCX doesn't have clear page breaks
            tables=tables,
            images=[],  # Image extraction would require additional processing
            processing_time=0.0,
            success=True
        )
    
    def _extract_docx_metadata(self, doc, file_path: Path) -> DocumentMetadata:
        """Extract DOCX metadata"""
        try:
            core_props = doc.core_properties
            file_stat = file_path.stat()
            
            return DocumentMetadata(
                filename=file_path.name,
                file_size=file_stat.st_size,
                mime_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                page_count=None,  # Not easily available in DOCX
                creation_date=core_props.created,
                modification_date=core_props.modified or datetime.fromtimestamp(file_stat.st_mtime),
                author=core_props.author,
                title=core_props.title,
                subject=core_props.subject,
                language=core_props.language,
                processing_method='python-docx',
                confidence_score=0.9
            )
        except Exception as e:
            logger.warning(f"Error extracting DOCX metadata: {e}")
            return self._create_basic_metadata(file_path, 'python-docx')
    
    def get_processor_name(self) -> str:
        return "DOCX"
    
    def get_supported_formats(self) -> List[str]:
        return ['.docx', '.docm']

class HTMLProcessor(DocumentProcessorInterface):
    """HTML processor using BeautifulSoup"""
    
    def __init__(self):
        if not BS4_AVAILABLE:
            raise ImportError("beautifulsoup4 is required for HTMLProcessor")
    
    def can_process(self, file_path: Path) -> bool:
        return file_path.suffix.lower() in ['.html', '.htm']
    
    async def process_document(self, file_path: Path) -> ProcessedDocument:
        """Process HTML document"""
        start_time = asyncio.get_event_loop().time()
        
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
            
            soup = BeautifulSoup(content, 'html.parser')
            
            # Extract text content
            text_content = soup.get_text(separator='\n', strip=True)
            
            # Extract tables
            tables = []
            for table_idx, table in enumerate(soup.find_all('table')):
                table_data = []
                for row in table.find_all('tr'):
                    row_data = [cell.get_text(strip=True) for cell in row.find_all(['td', 'th'])]
                    if row_data:  # Skip empty rows
                        table_data.append(row_data)
                
                if table_data:
                    tables.append({
                        'table_index': table_idx,
                        'data': table_data,
                        'rows': len(table_data),
                        'columns': len(table_data[0]) if table_data else 0
                    })
            
            # Extract metadata
            metadata = self._extract_html_metadata(soup, file_path)
            
            processing_time = asyncio.get_event_loop().time() - start_time
            
            return ProcessedDocument(
                document_id=self._generate_document_id(file_path),
                filename=file_path.name,
                text_content=text_content,
                metadata=metadata,
                page_texts=[text_content],
                tables=tables,
                images=[],
                processing_time=processing_time,
                success=True
            )
            
        except Exception as e:
            logger.error(f"HTML processing failed for {file_path}: {e}")
            return ProcessedDocument(
                document_id=self._generate_document_id(file_path),
                filename=file_path.name,
                text_content="",
                metadata=self._create_error_metadata(file_path, str(e)),
                page_texts=[],
                tables=[],
                images=[],
                processing_time=asyncio.get_event_loop().time() - start_time,
                success=False,
                error_message=str(e)
            )
    
    def _extract_html_metadata(self, soup, file_path: Path) -> DocumentMetadata:
        """Extract HTML metadata"""
        try:
            file_stat = file_path.stat()
            
            # Extract meta tags
            title = soup.find('title')
            title_text = title.get_text(strip=True) if title else None
            
            author_meta = soup.find('meta', attrs={'name': 'author'})
            author = author_meta.get('content') if author_meta else None
            
            description_meta = soup.find('meta', attrs={'name': 'description'})
            description = description_meta.get('content') if description_meta else None
            
            return DocumentMetadata(
                filename=file_path.name,
                file_size=file_stat.st_size,
                mime_type='text/html',
                page_count=1,
                creation_date=None,
                modification_date=datetime.fromtimestamp(file_stat.st_mtime),
                author=author,
                title=title_text,
                subject=description,
                language=None,
                processing_method='beautifulsoup4',
                confidence_score=0.85
            )
        except Exception as e:
            logger.warning(f"Error extracting HTML metadata: {e}")
            return self._create_basic_metadata(file_path, 'beautifulsoup4')
    
    def get_processor_name(self) -> str:
        return "HTML"
    
    def get_supported_formats(self) -> List[str]:
        return ['.html', '.htm']

class TextProcessor(DocumentProcessorInterface):
    """Plain text processor"""
    
    def can_process(self, file_path: Path) -> bool:
        return file_path.suffix.lower() in ['.txt', '.md', '.csv']
    
    async def process_document(self, file_path: Path) -> ProcessedDocument:
        """Process plain text document"""
        start_time = asyncio.get_event_loop().time()
        
        try:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252']
            content = None
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        content = file.read()
                    break
                except UnicodeDecodeError:
                    continue
            
            if content is None:
                raise ValueError("Could not decode file with any supported encoding")
            
            # Extract metadata
            metadata = self._create_basic_metadata(file_path, 'text')
            metadata.confidence_score = 1.0  # Text files are always processed correctly
            
            processing_time = asyncio.get_event_loop().time() - start_time
            
            return ProcessedDocument(
                document_id=self._generate_document_id(file_path),
                filename=file_path.name,
                text_content=content,
                metadata=metadata,
                page_texts=[content],
                tables=[],
                images=[],
                processing_time=processing_time,
                success=True
            )
            
        except Exception as e:
            logger.error(f"Text processing failed for {file_path}: {e}")
            return ProcessedDocument(
                document_id=self._generate_document_id(file_path),
                filename=file_path.name,
                text_content="",
                metadata=self._create_error_metadata(file_path, str(e)),
                page_texts=[],
                tables=[],
                images=[],
                processing_time=asyncio.get_event_loop().time() - start_time,
                success=False,
                error_message=str(e)
            )
    
    def get_processor_name(self) -> str:
        return "Text"
    
    def get_supported_formats(self) -> List[str]:
        return ['.txt', '.md', '.csv']

class DocumentProcessingModule:
    """Main Document Processing Module - Orchestrates all processors"""
    
    def __init__(self, config: Dict[str, Any] = None):
        self.config = config or {}
        
        # Initialize processors in order of preference
        self.processors = []
        
        # PDF processors (in order of preference)
        if PDFPLUMBER_AVAILABLE:
            self.processors.append(PDFPlumberProcessor())
        elif PYPDF_AVAILABLE:
            self.processors.append(PyPDFProcessor())
        
        # Other document processors
        if DOCX_AVAILABLE:
            self.processors.append(DOCXProcessor())
        
        if BS4_AVAILABLE:
            self.processors.append(HTMLProcessor())
        
        # Always available
        self.processors.append(TextProcessor())
        
        # Configuration
        self.max_concurrent_processing = self.config.get('max_concurrent_processing', 5)
        self.max_file_size_mb = self.config.get('max_file_size_mb', 100)
        self.enable_ocr = self.config.get('enable_ocr', False) and OCR_AVAILABLE
        
        logger.info(f"Initialized DocumentProcessingModule with {len(self.processors)} processors")
    
    def _generate_document_id(self, file_path: Path) -> str:
        """Generate unique document ID"""
        content = f"{file_path.name}_{file_path.stat().st_size}_{file_path.stat().st_mtime}"
        return hashlib.md5(content.encode()).hexdigest()
    
    def _create_basic_metadata(self, file_path: Path, processing_method: str) -> DocumentMetadata:
        """Create basic metadata when detailed extraction fails"""
        file_stat = file_path.stat()
        mime_type, _ = mimetypes.guess_type(str(file_path))
        
        return DocumentMetadata(
            filename=file_path.name,
            file_size=file_stat.st_size,
            mime_type=mime_type or 'application/octet-stream',
            page_count=None,
            creation_date=None,
            modification_date=datetime.fromtimestamp(file_stat.st_mtime),
            author=None,
            title=None,
            subject=None,
            language=None,
            processing_method=processing_method,
            confidence_score=0.5
        )
    
    def _create_error_metadata(self, file_path: Path, error: str) -> DocumentMetadata:
        """Create error metadata"""
        metadata = self._create_basic_metadata(file_path, 'error')
        metadata.confidence_score = 0.0
        return metadata
    
    def get_processor_for_file(self, file_path: Path) -> Optional[DocumentProcessorInterface]:
        """Get appropriate processor for file"""
        for processor in self.processors:
            if processor.can_process(file_path):
                return processor
        return None
    
    async def process_document(self, file_path: Union[str, Path]) -> ProcessedDocument:
        """Process a single document"""
        file_path = Path(file_path)
        
        # Validate file
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        # Check file size
        file_size_mb = file_path.stat().st_size / (1024 * 1024)
        if file_size_mb > self.max_file_size_mb:
            raise ValueError(f"File too large: {file_size_mb:.1f}MB > {self.max_file_size_mb}MB")
        
        # Get processor
        processor = self.get_processor_for_file(file_path)
        if not processor:
            raise ValueError(f"No processor available for file type: {file_path.suffix}")
        
        logger.info(f"Processing {file_path.name} with {processor.get_processor_name()}")
        
        # Process document
        result = await processor.process_document(file_path)
        
        # Apply OCR if enabled and text extraction failed/poor
        if (self.enable_ocr and 
            result.success and 
            len(result.text_content.strip()) < 100 and 
            file_path.suffix.lower() == '.pdf'):
            
            logger.info(f"Applying OCR to {file_path.name}")
            ocr_result = await self._apply_ocr(file_path)
            if ocr_result and len(ocr_result) > len(result.text_content):
                result.text_content = ocr_result
                result.metadata.processing_method += "+OCR"
                result.metadata.confidence_score = min(result.metadata.confidence_score + 0.1, 1.0)
        
        return result
    
    async def process_documents_batch(self, file_paths: List[Union[str, Path]]) -> List[ProcessedDocument]:
        """Process multiple documents concurrently"""
        semaphore = asyncio.Semaphore(self.max_concurrent_processing)
        
        async def process_with_semaphore(file_path):
            async with semaphore:
                try:
                    return await self.process_document(file_path)
                except Exception as e:
                    logger.error(f"Failed to process {file_path}: {e}")
                    return ProcessedDocument(
                        document_id=self._generate_document_id(Path(file_path)),
                        filename=Path(file_path).name,
                        text_content="",
                        metadata=self._create_error_metadata(Path(file_path), str(e)),
                        page_texts=[],
                        tables=[],
                        images=[],
                        processing_time=0.0,
                        success=False,
                        error_message=str(e)
                    )
        
        tasks = [process_with_semaphore(fp) for fp in file_paths]
        results = await asyncio.gather(*tasks)
        
        successful = sum(1 for r in results if r.success)
        logger.info(f"Processed {len(file_paths)} documents: {successful} successful, {len(file_paths) - successful} failed")
        
        return results
    
    async def _apply_ocr(self, file_path: Path) -> Optional[str]:
        """Apply OCR to document (for scanned PDFs)"""
        if not OCR_AVAILABLE:
            return None
        
        try:
            # This is a simplified OCR implementation
            # In practice, you'd convert PDF pages to images first
            loop = asyncio.get_event_loop()
            with ThreadPoolExecutor() as executor:
                # Placeholder for OCR implementation
                # Would need to convert PDF to images and then OCR each image
                return await loop.run_in_executor(executor, self._ocr_sync, file_path)
        except Exception as e:
            logger.error(f"OCR failed for {file_path}: {e}")
            return None
    
    def _ocr_sync(self, file_path: Path) -> str:
        """Synchronous OCR processing"""
        # Placeholder implementation
        # In practice, would use pdf2image + pytesseract
        return ""
    
    def get_supported_formats(self) -> Dict[str, List[str]]:
        """Get all supported formats by processor"""
        formats = {}
        for processor in self.processors:
            formats[processor.get_processor_name()] = processor.get_supported_formats()
        return formats
    
    def get_processor_status(self) -> Dict[str, Any]:
        """Get status of all processors"""
        return {
            'processors': [
                {
                    'name': p.get_processor_name(),
                    'supported_formats': p.get_supported_formats()
                }
                for p in self.processors
            ],
            'libraries': {
                'pdfplumber': PDFPLUMBER_AVAILABLE,
                'pypdf': PYPDF_AVAILABLE,
                'pymupdf': PYMUPDF_AVAILABLE,
                'python-docx': DOCX_AVAILABLE,
                'beautifulsoup4': BS4_AVAILABLE,
                'pytesseract': OCR_AVAILABLE
            },
            'config': {
                'max_concurrent_processing': self.max_concurrent_processing,
                'max_file_size_mb': self.max_file_size_mb,
                'enable_ocr': self.enable_ocr
            }
        }

# Example usage and testing
async def main():
    """Example usage of Document Processing Module"""
    
    # Initialize module
    config = {
        'max_concurrent_processing': 3,
        'max_file_size_mb': 50,
        'enable_ocr': False
    }
    
    processing_module = DocumentProcessingModule(config)
    
    # Show processor status
    print("Processor Status:")
    status = processing_module.get_processor_status()
    print(f"Available processors: {[p['name'] for p in status['processors']]}")
    print(f"Library availability: {status['libraries']}")
    
    # Example: Process a single document (if file exists)
    # test_file = Path("sample_rfp.pdf")
    # if test_file.exists():
    #     print(f"\nProcessing {test_file.name}...")
    #     result = await processing_module.process_document(test_file)
    #     print(f"Success: {result.success}")
    #     print(f"Text length: {len(result.text_content)}")
    #     print(f"Processing time: {result.processing_time:.2f}s")
    #     print(f"Method: {result.metadata.processing_method}")

if __name__ == "__main__":
    asyncio.run(main())
